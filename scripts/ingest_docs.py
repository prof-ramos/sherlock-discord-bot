import argparse
import asyncio
import logging
import re
import sys
from pathlib import Path

import tiktoken

# Ensure src is in python path
sys.path.append(str(Path(__file__).parent.parent))

from dotenv import load_dotenv

load_dotenv()

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None  # type: ignore

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None  # type: ignore

from src.rag_service import rag_service  # noqa: E402

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)


LEGAL_AUTHORITIES = [
    "stf",
    "stj",
    "tse",
    "tnu",
    "tcu",
    "agu",
    "tjrj",
    "tjsp",
    "tjmg",
    "tjrf",
    "trf",
    "trt",
    "tst",
    "stm",
    "anatel",
    "aneel",
    "antt",
    "anvisa",
]


def extract_legal_metadata(filename: str) -> dict:
    """
    Extracts legal metadata (Type, Number, Year, Authority) from filename.
    Examples:
    - "AGU (LC 73 de 1993).docx" -> type=LC, number=73, year=1993
    - "Súmulas STF atualizado.docx" -> type=Súmula, authority=STF
    - "Lei 8666.docx" -> type=Lei, number=8666
    """
    meta = {}
    clean_name = filename.lower()

    # Normalize simplified
    # (In a production env, we'd use unicodedata to remove accents, but simple replacement works for now)
    # Also replace separators with spaces to handle LC_73, Lei-8666, dec.123 correctly with word boundaries
    norm_name = clean_name.replace("ú", "u").replace("ç", "c").replace("ã", "a").replace("õ", "o")
    norm_name = norm_name.replace("_", " ").replace("-", " ").replace(".", " ")

    # 1. Detect Type (Order matters: specific -> general)
    # Using regex with word boundaries avoids matching inside other words
    patterns = [
        (r"\bsumulas?\b", "Súmula"),
        (r"\blei complementar\b", "Lei Complementar"),
        (r"\blc\b", "Lei Complementar"),
        (r"\bdecreto\s*lei\b", "Decreto-Lei"),
        (r"\bdl\b", "Decreto-Lei"),
        (r"\bdecreto\b", "Decreto"),
        (r"\bdec\b", "Decreto"),
        (r"\bportarias?\b", "Portaria"),
        (r"\bresolu[cç][aã]o\b", "Resolução"),
        (r"\bleis?\b", "Lei"),
    ]

    for pattern, type_name in patterns:
        if re.search(pattern, norm_name):
            meta["type"] = type_name
            break

    # 2. Detect Authority (Orgão)
    for auth in LEGAL_AUTHORITIES:
        # \b ensures we match "stf" but not "costfire"
        if re.search(r"\b" + re.escape(auth) + r"\b", norm_name):
            meta["authority"] = auth.upper()
            break

    # 3. Detect Year (19xx or 20xx)
    year_match = re.search(r"\b(19|20)\d{2}\b", clean_name)
    if year_match:
        meta["year"] = int(year_match.group())

    # 4. Detect Number
    # Heuristic: verify keywords first (nº, num, lei NNN, dec NNN)
    # Look for digits that are NOT the captured year

    # Prioritized search near keywords
    number_keywords = r"(?:n[º°o]|num|numero|lei|dec|decreto|lc|dl|portaria)\s*(\d+)"
    keyword_match = re.search(number_keywords, norm_name)

    candidate_number = None
    if keyword_match:
        candidate_number = keyword_match.group(1)

    # Fallback: find any digit sequence
    if not candidate_number:
        # Find all digit sequences
        all_numbers = re.findall(r"\d+", clean_name)
        for num in all_numbers:
            # Skip if it matches the year we already found
            if meta.get("year") and str(meta["year"]) == num:
                continue
            # Logic: valid law numbers usually 1 to ~6 digits.
            # Avoid single digits unless they are clearly numbered (handled by keyword match above)
            if len(num) >= 1 and len(num) <= 8:
                candidate_number = num
                break

    # Only assign fallback number if we have some context (Type or Authority) OR it came from a strong keyword
    # This prevents "123456.pdf" from being tagged as number=123456 without any other info
    if candidate_number:
        if keyword_match or meta.get("type") or meta.get("authority"):
            meta["number"] = candidate_number

    return meta


def extract_text_from_pdf(pdf_path: Path) -> str:
    if not PdfReader:
        logger.error("pypdf not installed. Cannot read PDF files.")
        sys.exit(1)

    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from DOCX including paragraphs and tables."""
    if not Document:
        logger.error("python-docx not installed. Cannot read DOCX files.")
        sys.exit(1)

    doc = Document(docx_path)
    text_parts = []

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract tables (per python-docx best practices from Context7)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def extract_text_from_html(html_path: Path) -> str:
    if not BeautifulSoup:
        logger.error("beautifulsoup4 not installed. Cannot read HTML files.")
        sys.exit(1)

    with open(html_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")

    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.extract()

    # Get text
    text = soup.get_text()

    # Clean up whitespace
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    text = "\n".join(chunk for chunk in chunks if chunk)

    return text


class RecursiveTokenSplitter:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50, model_name: str = "gpt-4o"):
        if chunk_overlap < 0:
            raise ValueError("chunk_overlap must be >= 0")
        if chunk_overlap >= chunk_size:
            raise ValueError(
                f"chunk_overlap ({chunk_overlap}) must be less than chunk_size ({chunk_size})"
            )

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tiktoken.encoding_for_model(model_name)
        self.separators = ["\n\n", "\n", ".", " ", ""]

    def split_text(self, text: str) -> list[str]:
        return self._split_text_recursive(text, self.separators)

    def _split_text_recursive(self, text: str, separators: list[str]) -> list[str]:
        separator = separators[-1]
        new_separators = []

        # Find the most appropriate separator
        for i, sep in enumerate(separators):
            if sep == "":
                separator = ""
                break
            if sep in text:
                separator = sep
                new_separators = separators[i + 1 :]
                break

        splits = text.split(separator) if separator else list(text)
        good_splits = []

        for split in splits:
            if not split.strip():
                continue

            # Reattach separator if it's not empty/whitespace
            if separator:
                split = (
                    split + separator
                )  # This is a simplification; ideally we keep separators intelligently

            if self._length(split) < self.chunk_size:
                good_splits.append(split)
            else:
                if new_separators:
                    good_splits.extend(self._split_text_recursive(split, new_separators))
                else:
                    # Forced split if no separators left
                    good_splits.extend(self._force_split(split))

        return self._merge_splits(good_splits)

    def _force_split(self, text: str) -> list[str]:
        """Split text that forces checking against token limit directly."""
        tokens = self.tokenizer.encode(text)
        chunks = []
        for i in range(0, len(tokens), self.chunk_size - self.chunk_overlap):
            chunk_tokens = tokens[i : i + self.chunk_size]
            chunks.append(self.tokenizer.decode(chunk_tokens))
        return chunks

    def _merge_splits(self, splits: list[str]) -> list[str]:
        """Merge smaller splits into chunks of max_size with overlap."""
        chunks = []
        current_chunk = []
        current_len = 0

        for split in splits:
            split_len = self._length(split)
            if current_len + split_len > self.chunk_size:
                if current_chunk:
                    joined_doc = "".join(current_chunk).strip()
                    if joined_doc:
                        chunks.append(joined_doc)

                        # Apply overlap: retain end of previous chunk for context
                        if self.chunk_overlap > 0 and len(joined_doc) > self.chunk_overlap:
                            overlap_text = joined_doc[-self.chunk_overlap :]
                            current_chunk = [overlap_text, split]
                            current_len = self._length(overlap_text) + split_len
                        else:
                            current_chunk = [split]
                            current_len = split_len
                    else:
                        current_chunk = [split]
                        current_len = split_len
                else:
                    current_chunk = [split]
                    current_len = split_len
            else:
                current_chunk.append(split)
                current_len += split_len

        if current_chunk:
            doc = "".join(current_chunk).strip()
            if doc:
                chunks.append(doc)

        return chunks

    def _length(self, text: str) -> int:
        return len(self.tokenizer.encode(text))


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Uses RecursiveTokenSplitter for chunking."""
    splitter = RecursiveTokenSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    return splitter.split_text(text)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".txt", ".md", ".text"}


def get_files_from_path(path: Path, recursive: bool = True) -> list[Path]:
    """Get all supported files from a path (file or directory)."""
    if path.is_file():
        if path.suffix.lower() in SUPPORTED_EXTENSIONS:
            return [path]
        return []

    # It's a directory
    if recursive:
        files = [
            f for f in path.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
        ]
    else:
        files = [
            f for f in path.iterdir() if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
        ]

    return sorted(files)


async def process_single_file(file_path: Path, chunk_size: int, overlap: int) -> bool:
    """Process a single file and add to RAG. Returns True on success."""
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            text = extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            text = extract_text_from_docx(file_path)
        elif suffix == ".html":
            text = extract_text_from_html(file_path)
        elif suffix in [".txt", ".md", ".text"]:
            text = file_path.read_text(encoding="utf-8")
        else:
            logger.warning("Unsupported format: %s", suffix)
            return False
    except Exception as e:
        logger.error("Failed to extract text from %s: %s", file_path.name, e)
        return False

    if not text.strip():
        logger.warning("No text found in %s", file_path.name)
        return False

    # Chunk text
    chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    logger.info("  🧩 Split into %d chunks", len(chunks))

    # Prepare data for RAG
    file_metadata = extract_legal_metadata(file_path.name)

    metadatas = []
    for i in range(len(chunks)):
        chunk_meta = {
            "source": file_path.name,
            "chunk_index": i,
            **file_metadata,  # Merge file-level metadata
        }
        metadatas.append(chunk_meta)

    try:
        if await rag_service.add_documents(chunks, metadatas):
            return True
        else:
            logger.error("  ❌ Failed to add to RAG")
            return False
    except Exception as e:
        logger.error("  ❌ RAG error: %s", e)
        return False


async def main():
    parser = argparse.ArgumentParser(description="Ingest documents into Sherlock's Knowledge Base")
    parser.add_argument("path", type=Path, help="Path to file or directory to ingest")
    parser.add_argument("--chunk-size", type=int, default=1000, help="Size of text chunks (tokens)")
    parser.add_argument("--overlap", type=int, default=200, help="Overlap between chunks (tokens)")
    parser.add_argument(
        "--no-recursive",
        action="store_true",
        help="Don't search subdirectories (only applies to directory input)",
    )

    args = parser.parse_args()
    input_path: Path = args.path.resolve()

    if not input_path.exists():
        logger.error("Path not found: %s", input_path)
        sys.exit(1)

    # Get list of files to process
    files = get_files_from_path(input_path, recursive=not args.no_recursive)

    if not files:
        logger.error("No supported files found. Supported: %s", ", ".join(SUPPORTED_EXTENSIONS))
        sys.exit(1)

    logger.info("📁 Found %d file(s) to process", len(files))

    # Process each file
    success_count = 0
    fail_count = 0

    for i, file_path in enumerate(files, 1):
        logger.info("[%d/%d] 📄 Processing: %s", i, len(files), file_path.name)

        if await process_single_file(file_path, args.chunk_size, args.overlap):
            logger.info("  ✅ Success")
            success_count += 1
        else:
            fail_count += 1

    # Summary
    logger.info("=" * 50)
    logger.info("📊 Ingestion Complete: %d success, %d failed", success_count, fail_count)

    if fail_count > 0:
        sys.exit(1)


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("Ingestion cancelled.")
    except Exception as e:
        logger.exception("Ingestion failed: %s", e)
